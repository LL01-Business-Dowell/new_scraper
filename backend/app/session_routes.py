"""
session_routes.py
-----------------
Handles the CSV chat session flow with file upload and multi-format output.

Endpoints:
  POST /sessions/
      Create new session from processed CSV.
  GET  /sessions/
      List all sessions (summary only, no CSV rows or messages).
  GET  /sessions/{session_id}
      Load full session including CSV data and message history.
  POST /sessions/{session_id}/message
      Send a message with optional file attachment.
      File content is extracted and sent to AI as reference context.
      AI response includes a reply text, optional CSV update, and output_type.
  GET  /sessions/{session_id}/download-docx/{msg_index}
      Generate and stream a .docx file from an assistant message reply.

File types supported for upload:
  .csv  — parsed as tabular data
  .xlsx — first sheet extracted as table
  .pdf  — text extracted page by page (requires pypdf)
  .docx — paragraph text extracted (requires python-docx)
  .txt  — read as plain text

Add to requirements.txt:
  pypdf
  python-docx
  openpyxl

Register in main.py:
  from .session_routes import router as session_router
  app.include_router(session_router)
"""

from fastapi import APIRouter, File, Form, UploadFile
from fastapi.responses import JSONResponse, FileResponse
from typing import List, Optional
import uuid
import datetime
import logging
import os
import io
import json
import re
import requests
import traceback
import pandas as pd

logger = logging.getLogger(__name__)
router = APIRouter()

# ---------------------------------------------------------------------------
# CRUD / Datacube — replace placeholder values before deploying
# ---------------------------------------------------------------------------
CRUD_BASE_URL    = os.getenv("CRUD_BASE_URL",           "https://datacube.uxlivinglab.online")
CRUD_API_KEY     = os.getenv("CRUD_API_KEY",            "sk_test_placeholder_replace_me")
CRUD_DATABASE_ID = os.getenv("DATABASE_ID",             "placeholder_database_id")
CRUD_COLLECTION  = os.getenv("SESSION_COLLECTION_NAME", "csv_chat_sessions")
CRUD_ENDPOINT    = f"{CRUD_BASE_URL.rstrip('/')}/api/crud"
CRUD_HEADERS     = {
    "Content-Type":  "application/json",
    "Authorization": f"Api-Key {CRUD_API_KEY}",
}

TMP_DIR = "/tmp"

# ---------------------------------------------------------------------------
# Shared Gemini rotator from main.py
# ---------------------------------------------------------------------------
try:
    from .main import gemini_rotator
    logger.info("session_routes: gemini_rotator imported from main")
except Exception as exc:
    logger.warning(f"session_routes: could not import gemini_rotator — {exc}")
    gemini_rotator = None


# ---------------------------------------------------------------------------
# File content extraction
# ---------------------------------------------------------------------------

def _extract_file_content(filename: str, file_bytes: bytes) -> str:
    """
    Extract readable text or tabular data from an uploaded file.
    Returns a plain-text string for injection into the Gemini prompt.
    Never raises — returns an error message string on failure.
    """
    ext = os.path.splitext(filename.lower())[1]
    logger.info(f"Extracting content: {filename} ({len(file_bytes)} bytes, type={ext})")

    try:
        if ext == ".csv":
            df = pd.read_csv(io.BytesIO(file_bytes))
            return (
                f"[Uploaded CSV — {len(df)} rows, {len(df.columns)} columns]\n"
                f"Columns: {', '.join(df.columns.tolist())}\n\n"
                f"{df.to_csv(index=False)}"
            )

        if ext in (".xlsx", ".xls"):
            df = pd.read_excel(io.BytesIO(file_bytes))
            return (
                f"[Uploaded Excel — {len(df)} rows, {len(df.columns)} columns]\n"
                f"Columns: {', '.join(df.columns.tolist())}\n\n"
                f"{df.to_csv(index=False)}"
            )

        if ext == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError:
                return "[PDF upload: pypdf is not installed — add pypdf to requirements.txt]"
            reader = PdfReader(io.BytesIO(file_bytes))
            pages  = []
            for i, page in enumerate(reader.pages):
                text = (page.extract_text() or "").strip()
                if text:
                    pages.append(f"--- Page {i+1} ---\n{text}")
            content = "\n\n".join(pages) if pages else "(no extractable text)"
            return f"[Uploaded PDF — {len(reader.pages)} pages]\n\n{content}"

        if ext == ".docx":
            try:
                from docx import Document
            except ImportError:
                return "[DOCX upload: python-docx is not installed — add python-docx to requirements.txt]"
            doc   = Document(io.BytesIO(file_bytes))
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            content = "\n\n".join(paras) if paras else "(no text found)"
            return f"[Uploaded Word document — {len(paras)} paragraphs]\n\n{content}"

        if ext == ".txt":
            return f"[Uploaded text file]\n\n{file_bytes.decode('utf-8', errors='replace')}"

        return (
            f"[File '{filename}' uploaded but format '{ext}' is not supported "
            f"for text extraction. Supported: .csv .xlsx .pdf .docx .txt]"
        )

    except Exception:
        logger.error(f"File extraction error for {filename}: {traceback.format_exc()}")
        return f"[Could not extract content from '{filename}' — check server logs]"


# ---------------------------------------------------------------------------
# DOCX generation from reply text
# ---------------------------------------------------------------------------

def _generate_docx(content: str, session_id: str, msg_index: int) -> str:
    """
    Write content into a .docx file.
    Lines starting with #/##/### become headings; others become paragraphs.
    Returns the file path.
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Add python-docx to requirements.txt"
        )

    doc = Document()
    doc.add_heading("Report", level=0)

    for block in content.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("### "):
            doc.add_heading(block[4:], level=3)
        elif block.startswith("## "):
            doc.add_heading(block[3:], level=2)
        elif block.startswith("# "):
            doc.add_heading(block[2:], level=1)
        else:
            # Handle inline markdown: strip ** for bold (Word styles handle that)
            doc.add_paragraph(block)

    path = os.path.join(TMP_DIR, f"session_{session_id}_msg{msg_index}.docx")
    doc.save(path)
    logger.info(f"DOCX written: {path}")
    return path


# ---------------------------------------------------------------------------
# CRUD helpers
# ---------------------------------------------------------------------------

def _crud_create(document: dict) -> bool:
    payload = {
        "database_id":     CRUD_DATABASE_ID,
        "collection_name": CRUD_COLLECTION,
        "documents":       [document],
    }
    try:
        resp = requests.post(CRUD_ENDPOINT, json=payload, headers=CRUD_HEADERS, timeout=30)
        ok   = resp.status_code in (200, 201)
        if not ok:
            logger.error(f"CRUD create failed {resp.status_code}: {resp.text[:300]}")
        return ok
    except Exception:
        logger.error(f"CRUD create exception: {traceback.format_exc()}")
        return False


def _crud_update(session_id: str, update_data: dict) -> bool:
    payload = {
        "database_id":     CRUD_DATABASE_ID,
        "collection_name": CRUD_COLLECTION,
        "filters":         {"session_id": session_id},
        "update_data":     update_data,
    }
    try:
        resp = requests.put(CRUD_ENDPOINT, json=payload, headers=CRUD_HEADERS, timeout=30)
        ok   = resp.status_code == 200
        if not ok:
            logger.error(f"CRUD update failed {resp.status_code}: {resp.text[:300]}")
        return ok
    except Exception:
        logger.error(f"CRUD update exception: {traceback.format_exc()}")
        return False


def _crud_get_one(session_id: str) -> Optional[dict]:
    params = {
        "database_id":     CRUD_DATABASE_ID,
        "collection_name": CRUD_COLLECTION,
        "filters":         json.dumps({"session_id": session_id}),
        "page":            1,
        "page_size":       1,
    }
    try:
        resp = requests.get(CRUD_ENDPOINT, params=params, headers=CRUD_HEADERS, timeout=30)
        if resp.status_code == 200:
            docs = resp.json().get("data", [])
            return docs[0] if docs else None
        logger.warning(f"CRUD get_one {resp.status_code} for {session_id}")
        return None
    except Exception:
        logger.error(f"CRUD get_one exception: {traceback.format_exc()}")
        return None


def _crud_list_all() -> List[dict]:
    params = {
        "database_id":     CRUD_DATABASE_ID,
        "collection_name": CRUD_COLLECTION,
        "filters":         "{}",
        "page":            1,
        "page_size":       100,
    }
    try:
        resp = requests.get(CRUD_ENDPOINT, params=params, headers=CRUD_HEADERS, timeout=30)
        if resp.status_code == 200:
            docs = resp.json().get("data", [])
            logger.info(f"CRUD list_all: {len(docs)} sessions")
            return docs
        logger.error(f"CRUD list_all failed {resp.status_code}: {resp.text[:300]}")
        return []
    except Exception:
        logger.error(f"CRUD list_all exception: {traceback.format_exc()}")
        return []


# ---------------------------------------------------------------------------
# Gemini prompt builder
# ---------------------------------------------------------------------------

def _build_session_prompt(
    user_prompt:      str,
    csv_columns:      List[str],
    csv_rows:         List[dict],
    message_history:  List[dict],
    attached_content: Optional[str],
) -> str:
    """
    Build the full prompt for a session exchange.

    output_type in the response tells the frontend what to offer:
      "csv"  — only CSV changed, show CSV download
      "text" — report/analysis produced, show DOCX download
      "both" — CSV changed AND substantial text produced
    """
    if csv_rows and csv_columns:
        df      = pd.DataFrame(csv_rows, columns=csv_columns)
        csv_str = df.to_csv(index=False)
    else:
        csv_str = "(no CSV data in this session)"

    history_lines = []
    for msg in message_history[-10:]:
        role = "User" if msg["role"] == "user" else "Assistant"
        history_lines.append(f"{role}: {msg['content'][:400]}")
    history_str = "\n".join(history_lines) if history_lines else "(no prior messages)"

    file_section = (
        f"\nATTACHED FILE CONTENT\n"
        f"---------------------\n"
        f"{attached_content}\n"
        if attached_content else ""
    )

    return f"""You are a flexible data and research assistant.
You help users analyse CSV datasets, process uploaded documents, and produce reports in any format.

CURRENT SESSION CSV DATA
------------------------
{csv_str}
{file_section}
CONVERSATION HISTORY (last 10 messages)
----------------------------------------
{history_str}

USER REQUEST
------------
{user_prompt}

RESPONSE FORMAT — FOLLOW EXACTLY
----------------------------------
Respond with ONLY a valid JSON object. Nothing before or after it.

{{
  "reply": "Full response text. For reports, write the entire report here. Markdown is supported (## headings, **bold**, bullet lists, tables). Can be as long as needed.",
  "csv_updated": true or false,
  "csv_data": [{{"col1": "val1", "col2": "val2"}}] or null,
  "output_type": "text" or "csv" or "both"
}}

Rules:
- reply: Always required. Write the complete response, report, or explanation here.
- csv_updated: true ONLY when modifying/filtering/adding to the CSV data.
- csv_data: Full updated dataset when csv_updated is true. Same column names as input. Null otherwise.
- output_type:
    "csv"  — only updated the table, no significant text output
    "text" — wrote a report, summary, or analysis (csv_updated must be false)
    "both" — updated CSV AND produced a substantial text response
- Never invent data not present in the provided context.
- First char must be {{ and last must be }}.
"""


# ---------------------------------------------------------------------------
# Gemini call + parser
# ---------------------------------------------------------------------------

def _call_gemini_for_session(
    user_prompt:      str,
    csv_columns:      List[str],
    csv_rows:         List[dict],
    message_history:  List[dict],
    session_id:       str,
    attached_content: Optional[str] = None,
) -> dict:
    result = {
        "reply":       "",
        "csv_updated": False,
        "csv_data":    None,
        "output_type": "text",
        "error":       None,
    }

    if gemini_rotator is None:
        result["error"] = "AI service is not available."
        result["reply"] = result["error"]
        return result

    prompt = _build_session_prompt(
        user_prompt, csv_columns, csv_rows, message_history, attached_content
    )
    logger.info(
        f"[SESSION {session_id}] Calling Gemini — "
        f"prompt={len(prompt)} chars, rows={len(csv_rows)}, file={'yes' if attached_content else 'no'}"
    )

    try:
        raw = gemini_rotator.call(prompt, temperature=0.0)
    except RuntimeError as exc:
        logger.error(f"[SESSION {session_id}] Rotator error: {exc}")
        result["error"] = str(exc)
        result["reply"] = str(exc)
        return result

    logger.info(f"[SESSION {session_id}] Response received ({len(raw)} chars)")

    # Strip markdown fences
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```[a-z]*\n?", "", cleaned)
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3].rstrip()

    # Extract outermost JSON object
    start = cleaned.find("{")
    end   = cleaned.rfind("}")
    if start == -1 or end == -1:
        logger.warning(f"[SESSION {session_id}] No JSON object found — using raw text as reply")
        result["reply"]       = cleaned or "No response received."
        result["output_type"] = "text"
        return result

    try:
        parsed = json.loads(cleaned[start:end + 1])
    except json.JSONDecodeError as exc:
        logger.error(f"[SESSION {session_id}] JSON parse failed: {exc}")
        result["reply"]       = cleaned
        result["output_type"] = "text"
        result["error"]       = f"Response parse issue (content still shown): {exc}"
        return result

    result["reply"]       = parsed.get("reply", "")
    result["output_type"] = parsed.get("output_type", "text")
    result["csv_updated"] = bool(parsed.get("csv_updated", False))

    csv_data = parsed.get("csv_data")
    if result["csv_updated"] and isinstance(csv_data, list) and len(csv_data) > 0:
        result["csv_data"] = csv_data
        logger.info(f"[SESSION {session_id}] CSV updated — {len(csv_data)} rows")
    else:
        result["csv_updated"] = False
        result["csv_data"]    = None
        if result["output_type"] == "csv":
            result["output_type"] = "text"

    return result


# ---------------------------------------------------------------------------
# API endpoints
# ---------------------------------------------------------------------------

@router.post("/sessions/")
async def create_session(
    source_task_id: str = Form(...),
    csv_columns:    str = Form(...),   # JSON-encoded list[str]
    csv_rows:       str = Form(...),   # JSON-encoded list[dict]
):
    """Create a new chat session from a processed CSV output."""
    try:
        columns = json.loads(csv_columns)
        rows    = json.loads(csv_rows)
    except json.JSONDecodeError as exc:
        return JSONResponse(
            status_code=400,
            content={"error": f"Invalid JSON in csv_columns or csv_rows: {exc}"},
        )

    session_id = str(uuid.uuid4())
    now        = datetime.datetime.utcnow().isoformat() + "Z"

    document = {
        "session_id":     session_id,
        "created_at":     now,
        "updated_at":     now,
        "source_task_id": source_task_id,
        "csv_columns":    columns,
        "csv_rows":       rows,
        "messages":       [],
        "title":          f"Session {datetime.datetime.utcnow().strftime('%b %d, %H:%M')}",
    }

    if not _crud_create(document):
        return JSONResponse(
            status_code=500,
            content={"error": "Failed to save session. Check CRUD configuration."},
        )

    logger.info(f"Session {session_id} created — {len(rows)} rows")
    return {"session_id": session_id, "title": document["title"], "created_at": now}


@router.get("/sessions/")
async def list_sessions():
    """Return all sessions ordered newest first as lightweight summaries."""
    docs = _crud_list_all()
    docs.sort(key=lambda d: d.get("created_at", ""), reverse=True)
    return {
        "sessions": [
            {
                "session_id":    d.get("session_id"),
                "title":         d.get("title", "Untitled"),
                "created_at":    d.get("created_at"),
                "updated_at":    d.get("updated_at"),
                "row_count":     len(d.get("csv_rows", [])),
                "message_count": len(d.get("messages", [])),
            }
            for d in docs
            if d.get("session_id")
        ]
    }


@router.get("/sessions/{session_id}")
async def get_session(session_id: str):
    """Load a full session document."""
    doc = _crud_get_one(session_id)
    if not doc:
        return JSONResponse(
            status_code=404,
            content={"error": f"Session '{session_id}' not found."},
        )
    doc.pop("_id", None)
    return doc


@router.post("/sessions/{session_id}/message")
async def send_message(
    session_id: str,
    prompt:     str                  = Form(...),
    file:       Optional[UploadFile] = File(None),
):
    """
    Send a message in a session, with an optional file attachment.

    The file's text content is extracted and provided to the AI as
    reference material alongside the existing CSV data and history.
    The session document in CRUD is overwritten after every exchange.
    """
    if not prompt.strip():
        return JSONResponse(
            status_code=400,
            content={"error": "Prompt must not be empty."},
        )

    doc = _crud_get_one(session_id)
    if not doc:
        return JSONResponse(
            status_code=404,
            content={"error": f"Session '{session_id}' not found."},
        )

    csv_columns = doc.get("csv_columns", [])
    csv_rows    = doc.get("csv_rows", [])
    messages    = doc.get("messages", [])

    # Extract uploaded file content
    attached_content  = None
    attached_filename = None
    if file and file.filename:
        file_bytes        = await file.read()
        attached_filename = file.filename
        attached_content  = _extract_file_content(file.filename, file_bytes)
        logger.info(
            f"[SESSION {session_id}] File: {file.filename} "
            f"({len(file_bytes)} bytes → {len(attached_content)} chars extracted)"
        )

    # Call AI
    ai_result = _call_gemini_for_session(
        user_prompt      = prompt.strip(),
        csv_columns      = csv_columns,
        csv_rows         = csv_rows,
        message_history  = messages,
        session_id       = session_id,
        attached_content = attached_content,
    )

    now = datetime.datetime.utcnow().isoformat() + "Z"

    user_msg = {
        "role":              "user",
        "content":           prompt.strip(),
        "timestamp":         now,
        "attached_filename": attached_filename,
    }
    assistant_msg = {
        "role":        "assistant",
        "content":     ai_result["reply"],
        "csv_updated": ai_result["csv_updated"],
        "output_type": ai_result["output_type"],
        "timestamp":   now,
        "error":       ai_result.get("error"),
    }

    messages.append(user_msg)
    messages.append(assistant_msg)
    assistant_msg_index = len(messages) - 1

    # Apply CSV update
    if ai_result["csv_updated"] and ai_result["csv_data"]:
        csv_rows    = ai_result["csv_data"]
        csv_columns = list(csv_rows[0].keys()) if csv_rows else csv_columns

    # Persist — overwrite session in CRUD
    _crud_update(session_id, {
        "updated_at":  now,
        "messages":    messages,
        "csv_rows":    csv_rows,
        "csv_columns": csv_columns,
    })

    return {
        "user_message":        user_msg,
        "assistant_message":   assistant_msg,
        "assistant_msg_index": assistant_msg_index,
        "csv_updated":         ai_result["csv_updated"],
        "csv_columns":         csv_columns,
        "csv_rows":            csv_rows,
        "output_type":         ai_result["output_type"],
        "error":               ai_result.get("error"),
    }


@router.get("/sessions/{session_id}/download-docx/{msg_index}")
def download_message_as_docx(session_id: str, msg_index: int):
    """
    Generate a .docx file from the reply text of a specific assistant message
    and stream it to the browser.
    """
    doc = _crud_get_one(session_id)
    if not doc:
        return JSONResponse(status_code=404, content={"error": "Session not found."})

    messages = doc.get("messages", [])
    if msg_index < 0 or msg_index >= len(messages):
        return JSONResponse(status_code=404, content={"error": "Message index out of range."})

    msg = messages[msg_index]
    if msg.get("role") != "assistant":
        return JSONResponse(
            status_code=400,
            content={"error": "Only assistant messages can be exported as DOCX."},
        )

    content = msg.get("content", "").strip()
    if not content:
        return JSONResponse(status_code=400, content={"error": "Message has no content."})

    try:
        docx_path = _generate_docx(content, session_id, msg_index)
    except RuntimeError as exc:
        return JSONResponse(status_code=500, content={"error": str(exc)})

    return FileResponse(
        path      = docx_path,
        filename  = f"report_{session_id[:8]}_msg{msg_index}.docx",
        media_type= "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )